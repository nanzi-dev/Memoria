"""Knowledge document extraction and structure-aware token chunking."""

from __future__ import annotations

from dataclasses import dataclass
from functools import lru_cache
from io import BytesIO
from pathlib import Path
import re
import stat
import zipfile

from memoria.core.config import configs


class KnowledgeDocumentError(ValueError):
    """Raised when an uploaded knowledge document cannot be safely processed."""


@dataclass(frozen=True)
class TextSection:
    text: str
    metadata: dict


@dataclass(frozen=True)
class ExtractedDocument:
    sections: list[TextSection]
    page_count: int | None = None

    @property
    def text(self) -> str:
        return "\n\n".join(section.text for section in self.sections if section.text)


_ALLOWED_SUFFIXES = {".txt", ".md", ".pdf", ".docx"}
_UNSUPPORTED_SUFFIXES = {".doc", ".xlsx", ".xls", ".pptx", ".ppt"}
_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MARKDOWN_IMAGE_RE = re.compile(
    r"^\s*(?:!\[\[[^\]]+\]\]|!\[[^\]]*\]\([^)]+\))\s*$"
)
_MARKDOWN_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-+:?\s*(?:\|\s*:?-+:?\s*)+\|?\s*$"
)
_SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[。！？!?；;])|(?<=\.)\s+|\n+")
_DOCX_MAX_ENTRIES = 2048
_DOCX_MAX_ENTRY_BYTES = 32 * 1024 * 1024
_DOCX_MAX_TOTAL_BYTES = 64 * 1024 * 1024
_DOCX_MAX_COMPRESSION_RATIO = 100


def validate_document_filename(filename: str) -> None:
    suffix = Path(filename or "").suffix.lower()
    if suffix in _UNSUPPORTED_SUFFIXES:
        raise KnowledgeDocumentError(f"不支持 {suffix or '该'} 格式")
    if suffix not in _ALLOWED_SUFFIXES:
        raise KnowledgeDocumentError("仅支持 TXT、Markdown、PDF 和 DOCX")


def validate_document_size(data: bytes) -> None:
    if not data:
        raise KnowledgeDocumentError("文档为空")
    if len(data) > configs.knowledge_upload_max_bytes:
        raise KnowledgeDocumentError("文档超过 10 MB 上传限制")


def extract_document(filename: str, data: bytes) -> ExtractedDocument:
    validate_document_size(data)
    validate_document_filename(filename)
    suffix = Path(filename or "").suffix.lower()

    if suffix == ".txt":
        extracted = _extract_utf8_text(data)
    elif suffix == ".md":
        extracted = _extract_markdown(data)
    elif suffix == ".pdf":
        extracted = _extract_pdf(data)
    else:
        extracted = _extract_docx(data)

    text = extracted.text.strip()
    if not text:
        raise KnowledgeDocumentError("文档未提取到可用文本")
    if len(text) > configs.knowledge_extract_max_chars:
        raise KnowledgeDocumentError("文档提取文本超过 1,000,000 字符限制")
    return extracted


def _extract_utf8_text(data: bytes) -> ExtractedDocument:
    if data.startswith((b"%PDF-", b"PK\x03\x04")) or b"\x00" in data[:4096]:
        raise KnowledgeDocumentError("文件内容与文本扩展名不匹配")
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise KnowledgeDocumentError("TXT/Markdown 必须使用 UTF-8 编码") from exc
    return ExtractedDocument([TextSection(_normalize_text(text), {})])


def _extract_markdown(data: bytes) -> ExtractedDocument:
    raw = _extract_utf8_text(data).text
    lines = raw.splitlines()
    if lines and lines[0].strip() == "---":
        for index in range(1, len(lines)):
            if lines[index].strip() == "---":
                lines = lines[index + 1 :]
                break

    sections: list[TextSection] = []
    heading_stack: list[str] = []
    table_index = 0
    index = 0

    def metadata(kind: str, **extra) -> dict:
        value = {"kind": kind, **extra}
        if heading_stack:
            value["heading_path"] = list(heading_stack)
        return value

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped or _MARKDOWN_IMAGE_RE.match(stripped):
            index += 1
            continue

        heading = _MARKDOWN_HEADING_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            title = _clean_markdown_inline(heading.group(2))
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(title)
            index += 1
            continue

        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
            block = [line]
            index += 1
            while index < len(lines):
                block.append(lines[index].rstrip())
                if lines[index].strip().startswith(fence):
                    index += 1
                    break
                index += 1
            sections.append(TextSection("\n".join(block), metadata("code")))
            continue

        if (
            "|" in stripped
            and index + 1 < len(lines)
            and _MARKDOWN_TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
        ):
            table_index += 1
            block = [
                _normalize_markdown_table_line(line),
                _normalize_markdown_table_separator(lines[index + 1]),
            ]
            index += 2
            while index < len(lines):
                candidate = lines[index].rstrip()
                if not candidate.strip() or "|" not in candidate:
                    break
                block.append(_normalize_markdown_table_line(candidate))
                index += 1
            sections.append(
                TextSection(
                    "\n".join(block),
                    metadata("table", table=table_index),
                )
            )
            continue

        if stripped.startswith(">"):
            block = []
            while index < len(lines) and lines[index].lstrip().startswith(">"):
                block.append(lines[index].rstrip())
                index += 1
            sections.append(TextSection("\n".join(block), metadata("quote")))
            continue

        block = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            candidate_stripped = candidate.strip()
            if (
                not candidate_stripped
                or _MARKDOWN_IMAGE_RE.match(candidate_stripped)
                or _MARKDOWN_HEADING_RE.match(candidate_stripped)
                or candidate_stripped.startswith(("```", "~~~", ">"))
            ):
                break
            if (
                "|" in candidate_stripped
                and index + 1 < len(lines)
                and _MARKDOWN_TABLE_SEPARATOR_RE.match(lines[index + 1].strip())
            ):
                break
            block.append(candidate)
            index += 1
        text = _normalize_text("\n".join(block))
        if text:
            sections.append(TextSection(text, metadata("paragraph")))

    if not sections and heading_stack:
        sections.append(
            TextSection(
                heading_stack[-1],
                {"kind": "heading", "heading_path": list(heading_stack)},
            )
        )
    return ExtractedDocument(sections)


def _clean_markdown_inline(text: str) -> str:
    text = re.sub(r"==(.+?)==", r"\1", text)
    text = re.sub(r"[*_`]+", "", text)
    return re.sub(r"\s+", " ", text).strip()


def _normalize_markdown_table_line(line: str) -> str:
    cells = [
        _clean_markdown_inline(cell)
        for cell in line.strip().strip("|").split("|")
    ]
    return "| " + " | ".join(cells) + " |"


def _normalize_markdown_table_separator(line: str) -> str:
    count = max(1, len(line.strip().strip("|").split("|")))
    return "| " + " | ".join("---" for _ in range(count)) + " |"


def _extract_pdf(data: bytes) -> ExtractedDocument:
    if not data.startswith(b"%PDF-"):
        raise KnowledgeDocumentError("文件内容与 PDF 扩展名不匹配")
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise KnowledgeDocumentError("PDF 文件损坏或无法读取") from exc
    if reader.is_encrypted:
        raise KnowledgeDocumentError("不支持加密 PDF")
    if len(reader.pages) > configs.knowledge_pdf_max_pages:
        raise KnowledgeDocumentError("PDF 超过 300 页限制")

    sections = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = _normalize_text(page.extract_text() or "")
        except Exception as exc:
            raise KnowledgeDocumentError(f"PDF 第 {page_number} 页解析失败") from exc
        if text:
            sections.append(TextSection(text, {"page": page_number}))
    if not sections:
        raise KnowledgeDocumentError("PDF 未包含可提取文本，暂不支持 OCR 扫描件")
    return ExtractedDocument(sections, page_count=len(reader.pages))


def _extract_docx(data: bytes) -> ExtractedDocument:
    if not data.startswith(b"PK\x03\x04") or not zipfile.is_zipfile(BytesIO(data)):
        raise KnowledgeDocumentError("文件内容与 DOCX 扩展名不匹配")
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > _DOCX_MAX_ENTRIES:
                raise KnowledgeDocumentError("DOCX 包含过多 ZIP 条目")
            total_size = 0
            names = set()
            for entry in entries:
                name = entry.filename
                path = Path(name.replace("\\", "/"))
                if (
                    not name
                    or "\x00" in name
                    or path.is_absolute()
                    or ".." in path.parts
                    or stat.S_ISLNK(entry.external_attr >> 16)
                ):
                    raise KnowledgeDocumentError("DOCX 包含不安全的 ZIP 条目")
                if entry.flag_bits & 0x1:
                    raise KnowledgeDocumentError("不支持加密 DOCX")
                if entry.file_size < 0 or entry.compress_size < 0:
                    raise KnowledgeDocumentError("DOCX ZIP 元数据无效")
                if entry.file_size > _DOCX_MAX_ENTRY_BYTES:
                    raise KnowledgeDocumentError("DOCX 单个 ZIP 条目解压后过大")
                total_size += entry.file_size
                if total_size > _DOCX_MAX_TOTAL_BYTES:
                    raise KnowledgeDocumentError("DOCX 解压后总大小超过限制")
                if entry.file_size:
                    if entry.compress_size == 0:
                        raise KnowledgeDocumentError("DOCX ZIP 压缩比异常")
                    if entry.file_size / entry.compress_size > _DOCX_MAX_COMPRESSION_RATIO:
                        raise KnowledgeDocumentError("DOCX ZIP 压缩比异常")
                names.add(name)
            if "word/document.xml" not in names:
                raise KnowledgeDocumentError("文件不是有效的 DOCX 文档")
        from docx import Document

        document = Document(BytesIO(data))
    except KnowledgeDocumentError:
        raise
    except Exception as exc:
        raise KnowledgeDocumentError("DOCX 文件损坏或无法读取") from exc

    sections = []
    for paragraph in document.paragraphs:
        text = _normalize_text(paragraph.text)
        if text:
            sections.append(TextSection(text, {"kind": "paragraph"}))
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [
                re.sub(r"\s+", " ", _normalize_text(cell.text))
                .replace("|", r"\|")
                .strip()
                for cell in row.cells
            ]
            rows.append(cells)
        if not rows or not any(any(cells) for cells in rows):
            continue
        formatted_rows = ["| " + " | ".join(cells) + " |" for cells in rows]
        separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
        row_count = len(rows)
        metadata = {"kind": "table", "table": table_index}
        if row_count == 1:
            metadata.update({"row_start": 1, "row_end": 1})
        else:
            metadata.update({"row_start": 2, "row_end": row_count})
        sections.append(
            TextSection(
                "\n".join([formatted_rows[0], separator, *formatted_rows[1:]]),
                metadata,
            )
        )
    if not sections:
        raise KnowledgeDocumentError("DOCX 未提取到可用段落或表格文本")
    return ExtractedDocument(sections)


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def chunk_document(
    extracted: ExtractedDocument,
    *,
    document_title: str = "",
    tokenizer=None,
    target_tokens: int | None = None,
    overlap_tokens: int | None = None,
    max_tokens: int | None = None,
    target_chars: int | None = None,
    overlap_chars: int | None = None,
) -> list[dict]:
    if target_chars is not None:
        codec = _CharacterCodec()
        target = max(1, target_chars)
        overlap = max(0, overlap_chars or 0)
        hard_limit = target
    else:
        codec = _TokenizerCodec(tokenizer or _get_default_tokenizer())
        target = target_tokens or configs.knowledge_chunk_target_tokens
        overlap = (
            configs.knowledge_chunk_overlap_tokens
            if overlap_tokens is None
            else overlap_tokens
        )
        hard_limit = max_tokens or configs.knowledge_chunk_max_tokens
    target = min(max(1, target), hard_limit)
    overlap = min(max(0, overlap), max(0, target - 1))

    chunks: list[dict] = []
    buffer: list[tuple[str, dict]] = []
    new_unit_count = 0
    active_heading: tuple[str, ...] | None = None

    def add_chunk(items: list[tuple[str, dict]]) -> None:
        content = "\n\n".join(text.strip() for text, _ in items if text.strip())
        if not content:
            return
        metadata = _merge_metadata([item[1] for item in items])
        index_text = build_chunk_index_text(document_title, metadata, content)
        if codec.count(index_text, special_tokens=True) > hard_limit:
            raise KnowledgeDocumentError(
                f"分块超过嵌入模型 {hard_limit} token 上限"
            )
        chunks.append(
            {
                "chunk_index": len(chunks),
                "content": content,
                "source_metadata": metadata,
            }
        )

    def flush(*, keep_overlap: bool) -> None:
        nonlocal buffer, new_unit_count
        if buffer and new_unit_count:
            add_chunk(buffer)
        if not keep_overlap or not buffer:
            buffer = []
            new_unit_count = 0
            return
        buffer = _overlap_tail(buffer, overlap, codec)
        new_unit_count = 0

    def append_unit(text: str, metadata: dict) -> None:
        nonlocal buffer, new_unit_count
        heading = tuple(metadata.get("heading_path") or ())
        if buffer and active_heading is not None and heading != active_heading:
            flush(keep_overlap=False)
        body_limit = _body_token_limit(
            codec,
            document_title,
            metadata,
            hard_limit,
        )
        target_body_limit = _body_token_limit(
            codec,
            document_title,
            metadata,
            target,
        )
        unit_limit = min(
            body_limit,
            max(1, target_body_limit - overlap),
        )
        for part in _split_text_unit(text, unit_limit, codec):
            candidate = [*buffer, (part, metadata)]
            if buffer and _indexed_count(
                candidate, document_title, codec
            ) > target:
                flush(keep_overlap=True)
                candidate = [*buffer, (part, metadata)]
            if buffer and _indexed_count(
                candidate, document_title, codec
            ) > hard_limit:
                buffer = []
                new_unit_count = 0
                candidate = [(part, metadata)]
            buffer = candidate
            new_unit_count += 1

    for section in extracted.sections:
        metadata = dict(section.metadata or {})
        heading = tuple(metadata.get("heading_path") or ())
        if active_heading is None:
            active_heading = heading
        elif heading != active_heading:
            flush(keep_overlap=False)
            active_heading = heading

        if metadata.get("kind") == "table":
            flush(keep_overlap=False)
            for item in _chunk_markdown_table(
                section.text,
                metadata,
                document_title,
                target,
                hard_limit,
                codec,
            ):
                add_chunk(item)
            active_heading = None
            continue

        isolated = metadata.get("kind") in {"code", "quote"}
        if isolated:
            flush(keep_overlap=False)
        paragraphs = [
            paragraph.strip()
            for paragraph in re.split(r"\n\s*\n", section.text)
            if paragraph.strip()
        ]
        for paragraph in paragraphs:
            for unit in _sentence_units(paragraph):
                append_unit(unit, metadata)
        if isolated:
            flush(keep_overlap=False)
            active_heading = None

    flush(keep_overlap=False)
    return chunks


def build_chunk_index_text(
    document_title: str,
    source_metadata: dict,
    content: str,
) -> str:
    parts = []
    title = Path(document_title or "").stem.strip()
    if title:
        parts.append(title)
    heading_path = source_metadata.get("heading_path") or []
    if isinstance(heading_path, str):
        heading_path = [heading_path]
    if heading_path:
        parts.append(" > ".join(str(item) for item in heading_path if item))
    parts.append(str(content or "").strip())
    return "\n".join(part for part in parts if part)


@lru_cache(maxsize=1)
def _get_default_tokenizer():
    from transformers import AutoTokenizer

    return AutoTokenizer.from_pretrained(configs.embedding_model)


class _TokenizerCodec:
    def __init__(self, tokenizer):
        self.tokenizer = tokenizer

    def encode(self, text: str) -> list:
        return list(
            self.tokenizer.encode(
                str(text or ""),
                add_special_tokens=False,
                truncation=False,
            )
        )

    def decode(self, tokens: list) -> str:
        return self.tokenizer.decode(tokens, skip_special_tokens=True).strip()

    def split(self, text: str, limit: int) -> list[str]:
        encoded = self.tokenizer(
            str(text or ""),
            add_special_tokens=False,
            truncation=False,
            return_offsets_mapping=True,
        )
        offsets = encoded.get("offset_mapping") or []
        if len(offsets) <= limit:
            return [str(text or "").strip()]
        parts = []
        start_token = 0
        while start_token < len(offsets):
            end_token = min(len(offsets), start_token + limit)
            start_char = offsets[start_token][0]
            end_char = offsets[end_token - 1][1]
            part = str(text or "")[start_char:end_char].strip()
            if part:
                parts.append(part)
            start_token = end_token
        return parts

    def tail(self, text: str, limit: int) -> str:
        encoded = self.tokenizer(
            str(text or ""),
            add_special_tokens=False,
            truncation=False,
            return_offsets_mapping=True,
        )
        offsets = encoded.get("offset_mapping") or []
        if len(offsets) <= limit:
            return str(text or "").strip()
        start_char = offsets[-limit][0]
        return str(text or "")[start_char:].strip()

    def count(self, text: str, *, special_tokens: bool = False) -> int:
        return len(
            self.tokenizer.encode(
                str(text or ""),
                add_special_tokens=special_tokens,
                truncation=False,
            )
        )


class _CharacterCodec:
    def encode(self, text: str) -> list[str]:
        return list(str(text or ""))

    def decode(self, tokens: list[str]) -> str:
        return "".join(tokens)

    def split(self, text: str, limit: int) -> list[str]:
        value = str(text or "")
        return [
            value[index : index + limit].strip()
            for index in range(0, len(value), limit)
            if value[index : index + limit].strip()
        ]

    def tail(self, text: str, limit: int) -> str:
        return str(text or "")[-limit:].strip()

    def count(self, text: str, *, special_tokens: bool = False) -> int:
        return len(str(text or ""))


def _sentence_units(text: str) -> list[str]:
    units = [
        unit.strip()
        for unit in _SENTENCE_BOUNDARY_RE.split(str(text or ""))
        if unit.strip()
    ]
    return units or [str(text or "").strip()]


def _split_text_unit(text: str, limit: int, codec) -> list[str]:
    return codec.split(text, limit)


def _body_token_limit(
    codec,
    document_title: str,
    metadata: dict,
    hard_limit: int,
) -> int:
    prefix = build_chunk_index_text(document_title, metadata, "")
    reserved = codec.count(prefix, special_tokens=True)
    return max(1, hard_limit - reserved - 1)


def _indexed_count(
    items: list[tuple[str, dict]],
    document_title: str,
    codec,
) -> int:
    content = "\n\n".join(text for text, _ in items)
    metadata = _merge_metadata([item[1] for item in items])
    return codec.count(
        build_chunk_index_text(document_title, metadata, content),
        special_tokens=True,
    )


def _overlap_tail(
    items: list[tuple[str, dict]],
    overlap: int,
    codec,
) -> list[tuple[str, dict]]:
    if overlap <= 0:
        return []
    selected: list[tuple[str, dict]] = []
    used = 0
    for text, metadata in reversed(items):
        count = codec.count(text)
        if selected and used + count > overlap:
            break
        if count > overlap:
            tail = codec.tail(text, overlap)
            return [(tail, metadata)] if tail else []
        selected.append((text, metadata))
        used += count
        if used >= overlap:
            break
    return list(reversed(selected))


def _chunk_markdown_table(
    text: str,
    metadata: dict,
    document_title: str,
    target: int,
    hard_limit: int,
    codec,
) -> list[list[tuple[str, dict]]]:
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    if len(lines) < 2:
        return [[(text, metadata)]]
    header = lines[:2]
    rows = lines[2:]
    if not rows:
        return [[("\n".join(header), metadata)]]

    result: list[list[tuple[str, dict]]] = []
    current_rows: list[str] = []
    first_row = max(1, int(metadata.get("row_start") or 1))
    row_start = first_row

    def table_items(selected_rows: list[str], start: int) -> list[tuple[str, dict]]:
        item_metadata = {
            **metadata,
            "row_start": start,
            "row_end": start + len(selected_rows) - 1,
        }
        return [("\n".join([*header, *selected_rows]), item_metadata)]

    for row_number, row in enumerate(rows, start=first_row):
        candidates = [*current_rows, row]
        items = table_items(candidates, row_start)
        count = _indexed_count(items, document_title, codec)
        if current_rows and count > target:
            result.append(table_items(current_rows, row_start))
            current_rows = []
            row_start = row_number
            items = table_items([row], row_start)
            count = _indexed_count(items, document_title, codec)
        if count > hard_limit:
            if current_rows:
                result.append(table_items(current_rows, row_start))
                current_rows = []
            fragments = _split_table_row(
                row,
                header,
                metadata,
                document_title,
                hard_limit,
                codec,
                row_number,
            )
            result.extend(fragments)
            row_start = row_number + 1
            continue
        current_rows.append(row)
    if current_rows:
        result.append(table_items(current_rows, row_start))
    return result


def _split_table_row(
    row: str,
    header: list[str],
    metadata: dict,
    document_title: str,
    hard_limit: int,
    codec,
    row_number: int,
) -> list[list[tuple[str, dict]]]:
    prefix = "\n".join(header)
    body_limit = _body_token_limit(codec, document_title, metadata, hard_limit)
    available = max(1, body_limit - codec.count(prefix) - 8)
    fragments = _split_text_unit(row.strip().strip("|"), available, codec)
    result = []
    for index, fragment in enumerate(fragments, start=1):
        continuation = f"| {fragment} |"
        item_metadata = {
            **metadata,
            "row_start": row_number,
            "row_end": row_number,
            "row_fragment": index,
        }
        result.append([("\n".join([*header, continuation]), item_metadata)])
    return result


def _merge_metadata(items: list[dict]) -> dict:
    pages = sorted({item["page"] for item in items if item.get("page") is not None})
    result = {}
    if pages:
        result["pages"] = pages
        if len(pages) == 1:
            result["page"] = pages[0]
    kinds = sorted({item["kind"] for item in items if item.get("kind")})
    if kinds:
        result["kinds"] = kinds
        if len(kinds) == 1:
            result["kind"] = kinds[0]
    heading_paths = {
        tuple(item.get("heading_path") or ())
        for item in items
        if item.get("heading_path")
    }
    if len(heading_paths) == 1:
        result["heading_path"] = list(next(iter(heading_paths)))
    tables = sorted({item["table"] for item in items if item.get("table") is not None})
    if len(tables) == 1:
        result["table"] = tables[0]
    rows = [item["row"] for item in items if item.get("row") is not None]
    row_starts = [
        item["row_start"]
        for item in items
        if item.get("row_start") is not None
    ]
    row_ends = [
        item["row_end"]
        for item in items
        if item.get("row_end") is not None
    ]
    if rows:
        row_starts.extend(rows)
        row_ends.extend(rows)
    if row_starts:
        result["row_start"] = min(row_starts)
    if row_ends:
        result["row_end"] = max(row_ends)
    if row_starts and row_ends and min(row_starts) == max(row_ends):
        result["row"] = min(row_starts)
    row_fragments = [
        item["row_fragment"]
        for item in items
        if item.get("row_fragment") is not None
    ]
    if len(row_fragments) == 1:
        result["row_fragment"] = row_fragments[0]
    return result
